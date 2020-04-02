#!/usr/bin/python3
# coding=utf-8

import sys
import getopt
import argparse
import docx
import os

def printhelp():
    print("Usage: ./SymDetection.py [-f|--file file] [options]")
    print("       ./SymDetection.py [-h|--help] [-v|--version]\n")
    print("-f --file      待检测.docx文件名")
    print("-o --origin    检测起始段落号，选填")
    print("-n --num       检测段落数，选填\n")
    print("-h --help      打印帮助信息")
    print("-v --version   打印版本号")

def checkfile(file_name, sta_parag, parag_num):
    file = docx.Document(file_name)
    if sta_parag < 1:
        sta_parag = 1
    elif sta_parag > len(file.paragraphs):
        print("\033[1;31m起始段落超出全文大小！\033[0m")
        return -1
    if parag_num < 1:
        parag_num = len(file.paragraphs)
    
    #检测符号数
    for i in range(len(file.paragraphs)):
        if ((i + 1) >= sta_parag) and ((i + 1) < sta_parag + parag_num):
            comma_num = file.paragraphs[i].text.count('，')
            period_num = file.paragraphs[i].text.count('。')
            if period_num != 1:
                print("第 %d 段：逗号数 %d，句号数 %d，\033[1;31m句号数超出标准！\033[0m\n\t%s\n" % (i + 1, comma_num, period_num, file.paragraphs[i].text))
            else:
                print("第 %d 段：逗号数 %d，句号数 %d" % (i + 1, comma_num, period_num))

if __name__ == '__main__':
    file_name = "null"
    origin_par = 0
    par_num = 0

    opts, args = getopt.getopt(sys.argv[1:], "-h-v-f:-o:-n:", ["help", "file=", "version", "origin=", "num="])
    for opt, arg in opts:
        if opt in ("-v", "--version"):
            print("version 1.0.0 (12 Aug 2019)")
            sys.exit()
        elif opt in ("-h", "--help"):
            printhelp()
            sys.exit()    
        elif opt in ("-f", "--file"):
            file_name = arg
        elif opt in ("-o", "--origin"):
            origin_par = int(arg)
        elif opt in ("-n", "--num"):
            par_num = int(arg)

    if file_name == "null":
        print("Usage: python3 SymDetection.py [-f|--file file] [options]")
        print("Try 'python3 SymDetection.py --help' for more information.")
        sys.exit()
    else:
        #linux下需要下载antiword把doc转化docx，暂时不兼容了
        if file_name.find('.docx') == -1:
            print("目前仅支持docx格式文件!\n")
        else:
            print("正在检测文件：>>>\033[1;31m%s\033[0m<<<\n" %  file_name)
            #检测模式
            checkfile(file_name, origin_par, par_num)

